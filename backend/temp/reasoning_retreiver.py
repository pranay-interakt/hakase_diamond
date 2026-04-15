"""
PageIndex-style Reasoning Retriever for Hakase Clinical AI.

Implements the core ideas from VectifyAI/PageIndex:
  - No chunking, no vector DB, no embeddings
  - Builds a hierarchical structural tree from the document (TOC + section headers)
  - Retrieves pages by reasoning about which tree branches are relevant to a query
  - Operates directly on page-level text, handling tables and scanned PDFs via OCR

Flow:
  1. Load document → list of {page_num, text, tables, headers}
  2. Build PageIndex tree → hierarchical sections {section_name: {parent: ..., pages: []}}
  3. On query → traverse tree using semantic synonyms, score pages, return top-k
"""

import logging
import re
import os
import difflib
from typing import List, Dict, Optional, Any

logger = logging.getLogger(__name__)

# ── Section taxonomy & Synonyms ───────────────────────────────────────────────
SECTION_ANCHORS: Dict[str, List[str]] = {
    "synopsis":       ["synopsis", "study summary", "protocol summary", "abstract", "executive summary"],
    "objectives":     ["objective", "primary objective", "secondary objective", "study goal", "purpose"],
    "endpoints":      ["endpoint", "outcome measure", "efficacy variable", "primary endpoint", "estimand"],
    "design":         ["study design", "trial design", "methodology", "randomisation", "randomization", "blinding", "open-label", "study plan"],
    "eligibility":    ["inclusion criteria", "exclusion criteria", "eligibility", "selection of subjects", "patient population", "criteria for evaluation"],
    "dosing":         ["dose", "dosage", "administration", "posology", "treatment schedule", "study drug preparation"],
    "schedule":       ["schedule of events", "assessment schedule", "visit schedule", "flowchart", "study calendar", "time and events"],
    "adverse_events": ["adverse event", "adr", "safety monitoring", "toxicity", "tolerability", "serious adverse"],
    "stopping_rules": ["stopping rule", "discontinuation", "withdrawal criteria", "early termination", "subject withdrawal"],
    "statistics":     ["statistical analysis", "sample size", "power calculation", "efficacy analysis", "biostatistics", "statistical methods"],
    "ethics":         ["ethics committee", "irb", "institutional review", "informed consent", "gcp", "good clinical practice"],
    "regulatory":     ["regulatory", "ema", "fda", "ich", "competent authority", "marketing authorisation"],
    "investigators":  ["investigator", "principal investigator", "coordinating investigator"],
    "sites":          ["study site", "clinical site", "site selection", "participating centre", "study location"],
}

# General synonyms for fuzzy matching / query expansion
SYNONYMS = {
    "drug": ["medication", "treatment", "intervention", "therapy", "agent", "product"],
    "disease": ["condition", "indication", "disorder", "illness", "syndrome"],
    "patient": ["subject", "participant", "volunteer"],
    "safety": ["adverse", "toxicity", "tolerability", "risk", "hazard"],
    "efficacy": ["effectiveness", "benefit", "outcome", "response"]
}


class PageIndex:
    """
    Hierarchical page-level index. Enhances raw text with hierarchical tracking
    and expanded semantic scoring rules.
    """

    def __init__(self):
        self.pages: List[Dict] = []          # [{page_num, text, tables, headings}]
        self.index: Dict[str, List[int]] = {}  # section → [page_nums]
        self.toc: List[Dict] = []            # [{level, title, page_num}]
        self.tree: Dict[str, Any] = {}       # True hierarchical tree representation

    def build(self, pages: List[Dict], toc: Optional[List[Dict]] = None):
        self.pages = pages
        if toc:
            self.toc = toc
            self._index_from_toc(toc)
        self._index_from_content()
        logger.info(f"[PageIndex] Built hierarchical index: {len(self.index)} structured sections, {len(self.pages)} pages")

    def _index_from_toc(self, toc: List[Dict]):
        """Build a hierarchical index prioritizing tree-level parents."""
        current_hier = []
        for entry in toc:
            level = entry.get("level", 1)
            title = entry.get("title", "").strip()
            page_num = entry.get("page_num", 0)
            
            # Maintain hierarchy tree
            while len(current_hier) >= level:
                current_hier.pop()
            current_hier.append(title)
            
            title_lower = title.lower()
            
            # Map to canonical anchors
            for section, anchors in SECTION_ANCHORS.items():
                if any(a in title_lower for a in anchors):
                    self.index.setdefault(section, [])
                    if page_num not in self.index[section]:
                        self.index[section].append(page_num)
                    # Create tree node
                    self.tree[section] = {
                        "hierarchy_path": list(current_hier),
                        "pages": self.index[section]
                    }

    def _index_from_content(self):
        """Scan page content, layout hints (headings), and tables."""
        for p in self.pages:
            text_lower = p["text"].lower()
            
            # Analyze detected headings (if provided by parser) and plain text
            headings_lower = [h.lower() for h in p.get("headings", [])]
            search_area = " ".join(headings_lower) + " " + text_lower[:1000]

            for section, anchors in SECTION_ANCHORS.items():
                for anchor in anchors:
                    if re.search(r'\b' + re.escape(anchor) + r's?\b', search_area):
                        self.index.setdefault(section, [])
                        pn = p["page_num"]
                        if pn not in self.index[section]:
                            self.index[section].append(pn)
                        break

    def get_pages_for_section(self, section: str, max_pages: int = 3) -> List[Dict]:
        page_nums = self.index.get(section, [])
        if not page_nums:
            page_nums = self._fallback_scan(section)
        
        results = []
        seen = set()
        for pn in page_nums[:max_pages]:
            if pn not in seen:
                page = next((p for p in self.pages if p["page_num"] == pn), None)
                if page:
                    results.append(page)
                    seen.add(pn)
        return results

    def query(self, query: str, top_k: int = 4) -> List[Dict]:
        """
        Enhanced retrieval using synonyms & semantic understanding, scoring
        regular text plus explicitly extracted table tabular data.
        """
        query_lower = query.lower()
        query_tokens = set(re.findall(r'\b\w{4,}\b', query_lower))
        
        # Query expansion using synonyms
        expanded_tokens = set(query_tokens)
        for token in query_tokens:
            for k, syns in SYNONYMS.items():
                if token == k or token in syns:
                    expanded_tokens.add(k)
                    expanded_tokens.update(syns)

        scores = []
        for p in self.pages:
            text_lower = p["text"].lower()
            page_tokens = set(re.findall(r'\b\w{4,}\b', text_lower))
            
            # Semantic overlap
            overlap = len(expanded_tokens & page_tokens)
            
            # Bonus if terms appear in headings
            header_bonus = 0
            for h in p.get("headings", []):
                h_tokens = set(re.findall(r'\b\w{4,}\b', h.lower()))
                header_bonus += len(expanded_tokens & h_tokens) * 3
                
            # Table data is highly critical: double overlap points if in tables
            tables_text = " ".join(p.get("tables", [])).lower()
            table_tokens = set(re.findall(r'\b\w{4,}\b', tables_text))
            table_bonus = len(expanded_tokens & table_tokens) * 2

            total_score = overlap + header_bonus + table_bonus
            scores.append((p["page_num"], total_score, p))

        scores.sort(key=lambda x: x[1], reverse=True)
        return [s[2] for s in scores[:top_k] if s[1] > 0]

    def _fallback_scan(self, section: str) -> List[int]:
        keyword = section.replace("_", " ")
        results = []
        for p in self.pages:
            if keyword in p["text"].lower()[:1500] or keyword in " ".join(p.get("headings", [])).lower():
                results.append(p["page_num"])
        if not results and section == "synopsis":
            results = [p["page_num"] for p in self.pages[:3]]
        return results

    def get_overview_context(self, max_chars: int = 24000) -> str:
        """
        Build a high-yield context string by pulling the top pages per priority
        section. Each section gets a capped budget so all sections are represented
        even for very large protocols. Deduplicates pages seen across sections.
        """
        priority = ["synopsis", "objectives", "endpoints", "design", "eligibility"]
        # Fair budget per section so no one section starves the rest
        per_section_budget = max_chars // max(len(priority), 1)
        parts = ["=== PROTOCOL PAGEINDEX OVERVIEW ===\n"]
        seen_pages: set = set()

        for section in priority:
            pages = self.get_pages_for_section(section, max_pages=3)
            section_chars = 0
            for page in pages:
                pn = page["page_num"]
                if pn in seen_pages:
                    continue
                seen_pages.add(pn)
                tbl_txt = ("\n[TABLES]\n" + "\n".join(page.get("tables", []))) if page.get("tables") else ""
                raw = f"\n[SECTION: {section.upper()} | PAGE {pn}]\n{page['text']}{tbl_txt}\n"
                # Truncate the individual snippet to the per-section budget
                allowed = per_section_budget - section_chars
                if allowed <= 0:
                    break
                snippet = raw[:allowed]
                parts.append(snippet)
                section_chars += len(snippet)

        result = "".join(parts)
        print(f"      [PageIndex Context Builder] Built {len(result)} chars of context from {len(seen_pages)} pages.")
        # Hard fallback: if nothing useful was built, return first pages verbatim
        if len(result) < 400:
            fallback_pages = self.pages[:6]
            fallback = "\n".join(
                f"[PAGE {p['page_num']}]\n{p['text']}" for p in fallback_pages
            )
            result = f"=== FALLBACK: FIRST PAGES ===\n{fallback[:max_chars]}"
        return result


class ReasoningRetriever:
    """
    Public interface for PageIndex retrieval.
    Includes OCR fallbacks, table extraction, and robust document parsing.
    """

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.page_index = PageIndex()
        self._load_and_build()

    def _apply_ocr_fallback(self, page_img) -> str:
        """OCR fallback for scanned PDFs."""
        try:
            import pytesseract
            return pytesseract.image_to_string(page_img)
        except Exception as e:
            logger.warning(f"OCR fallback failed or pytesseract not available: {e}")
            return ""

    def _load_and_build(self):
        pages = []
        toc = []
        ext = os.path.splitext(self.file_path)[1].lower()

        try:
            if ext == ".pdf":
                import fitz  # PyMuPDF
                doc = fitz.open(self.file_path)
                
                # Extract TOC (bookmarks) with levels
                raw_toc = doc.get_toc()  # [[level, title, page], ...]
                toc = [{"level": t[0], "title": t[1], "page_num": t[2]} for t in raw_toc]
                
                for i, page in enumerate(doc):
                    text = page.get_text("text")
                    
                    # 1. OCR Fallback for scanned/empty pages
                    if len(text.strip()) < 50:
                        pix = page.get_pixmap()
                        if pix:
                            from PIL import Image
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                            ocr_text = self._apply_ocr_fallback(img)
                            if ocr_text.strip():
                                text = ocr_text

                    # 2. Extract semantically meaningful layout-based headings
                    font_sizes = []
                    dict_blocks = page.get_text("dict", flags=11).get("blocks", [])
                    for b in dict_blocks:
                        for l in b.get("lines", []):
                            for s in l.get("spans", []):
                                font_sizes.append(s["size"])
                    
                    from collections import Counter
                    mode_size = Counter(font_sizes).most_common(1)[0][0] if font_sizes else 11

                    headings = []
                    for b in dict_blocks:
                        for l in b.get("lines", []):
                            for s in l.get("spans", []):
                                text_str = s["text"].strip()
                                if not text_str:
                                    continue
                                size = s["size"]
                                is_bold = "bold" in s["font"].lower()
                                
                                # Semantic heading logic: larger than body text, or short/bold/numbered
                                is_larger = (size > mode_size + 0.5) and len(text_str) < 100
                                is_bold_numbered = (abs(size - mode_size) <= 0.5) and is_bold and len(text_str) < 80 and (text_str.istitle() or re.match(r'^\d+(\.\d+)*\s+', text_str))
                                
                                if is_larger or is_bold_numbered:
                                    headings.append(text_str)

                    # 3. Explicit Table Extraction
                    tables_text = []
                    tabs = page.find_tables()
                    if tabs and tabs.tables:
                        for tab in tabs.tables:
                            for row in tab.extract():
                                tables_text.append(" | ".join([str(c) if c else "" for c in row]))

                    pages.append({
                        "page_num": i + 1,
                        "text": text,
                        "headings": headings,
                        "tables": tables_text
                    })
                doc.close()

            elif ext == ".docx":
                # Robust DOCX Parsing
                from docx import Document as DocxDocument
                doc = DocxDocument(self.file_path)
                
                page_counter = 1
                current_text = []
                current_headings = []
                
                for para in doc.paragraphs:
                    # Treat Heading styles as structural breaks/pages
                    if para.style.name.startswith('Heading'):
                        if current_text:
                            pages.append({
                                "page_num": page_counter,
                                "text": "\n".join(current_text),
                                "headings": current_headings,
                                "tables": []
                            })
                            page_counter += 1
                            current_text = []
                            current_headings = []
                        current_headings.append(para.text)
                    current_text.append(para.text)
                
                # Add final trailing content
                if current_text:
                    pages.append({
                        "page_num": page_counter,
                        "text": "\n".join(current_text),
                        "headings": current_headings,
                        "tables": []  # Requires deeper parsing for docx tables
                    })
            else:
                with open(self.file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                for i, start in enumerate(range(0, len(text), 3000)):
                    pages.append({"page_num": i + 1, "text": text[start:start + 3000], "headings": [], "tables": []})

        except Exception as e:
            logger.error(f"[ReasoningRetriever] Failed to load {self.file_path}: {e}")

        self.page_index.build(pages, toc or None)

    # ── Public API ────────────────────────────────────────────────────────────

    def get_section_context(self, section_name: str, max_pages: int = 3) -> str:
        pages = self.page_index.get_pages_for_section(section_name, max_pages=max_pages)
        if not pages:
            return ""
        parts = [f"\n[PAGE {p['page_num']} — {section_name.upper()}]\n{p['text']}" for p in pages]
        return "".join(parts)

    def get_high_yield_overview_context(self) -> str:
        return self.page_index.get_overview_context(max_chars=8000)

    def query(self, question: str, top_k: int = 4) -> str:
        pages = self.page_index.query(question, top_k=top_k)
        if not pages:
            return ""
        parts = [f"\n[PAGE {p['page_num']} — RETRIEVED FOR: {question[:60]}]\n{p['text']}" for p in pages]
        return "".join(parts)

    def agentic_explore(self, objective: str, llm) -> str:
        """
        Agent-driven document exploration loop replacing static queries.
        """
        agent = AgenticDocumentExplorer(self.page_index, llm, max_steps=8)
        return agent.run(objective)


class AgenticDocumentExplorer:
    """
    An agentic loop that dynamically decides which pages/sections to pull
    to satisfy a high-level extraction objective, instead of naive keyword search.
    """
    def __init__(self, page_index, llm, max_steps=5):
        self.page_index = page_index
        self.llm = llm
        self.max_steps = max_steps
        
    def run(self, objective: str) -> str:
        memory = []
        pages_seen = set()
        
        for step in range(self.max_steps):
            tools_desc = '''Available Tools:
1. SEARCH(query: str): Semantic search across the document. Returns top 3 page hints.
2. GET_PAGE(page_num: int): Retrieves full text and structured tables for a specific page.
3. GET_METADATA(): Retrieves the document hierarchical Table of Contents and detected sections.
'''
            context = "\n".join([f"Step {m['step']}: Action={m['action']}({m['arg']}) -> Result Summary={m['result'][:300]}..." for m in memory])
            
            prompt = f"""You are a Clinical Protocol Reasoning Agent. 
Objective: {objective}

{tools_desc}

Current Step: {step}/{self.max_steps}
Guidelines:
- Start with GET_METADATA() to understand the document structure.
- Use SEARCH(query) to find where specific topics (like "sample size" or "primary endpoint") are discussed.
- Use GET_PAGE(page_num) to read the full context once you have a potential page. 
- Do NOT repeat the same action if it didn't give you new information.
- If you find a number or specific field, VERIFY it against the surrounding text in your thought process.

Past Actions:
{context}

Decide your next action. Reply strictly in this format:
THOUGHT: <your reasoning for this specific step>
ACTION: <tool_name>
ARG: <argument_string>

If you have definitively found the answer, use:
ACTION: FINAL_ANSWER
ARG: <the final structured answer or specific data requested>
"""
            print(f"    [Agent Step {step}] Thinking...")
            try:
                response = _fallback_generate_agent(self.llm, prompt)
            except Exception as e:
                print(f"    [Agent Step {step}] LLM Error: {e}")
                return f"Error connecting to LLM reasoning: {e}"
            
            action_match = re.search(r'ACTION:\s*([A-Z_]+)', response)
            arg_match = re.search(r'ARG:\s*(.+)', response)
            thought_match = re.search(r'THOUGHT:\s*(.+)', response, re.DOTALL | re.IGNORECASE)
            
            thought = thought_match.group(1).split('ACTION:')[0].strip() if thought_match else "No thought provided"
            
            if not action_match or not arg_match:
                print(f"    [Agent Step {step}] Error: Could not parse response. Response output: {response[:100]}...")
                memory.append({"step": step, "action": "ERROR", "arg": "", "result": "Failed to parse action from response. Ensure you use ACTION: and ARG: format."})
                continue
                
            action = action_match.group(1).strip()
            arg = arg_match.group(1).strip()
            
            print(f"    [Agent Step {step}] Action: {action}({arg})")
            print(f"    [Agent Step {step}] Thought: {thought[:150]}...")
            
            if action == "FINAL_ANSWER":
                return self._verify_and_return(objective, arg)
                
            elif action == "SEARCH":
                results = self.page_index.query(arg, top_k=3)
                hints = []
                for r in results:
                    clean_text = r['text'][:300].replace('\n', ' ')
                    hints.append(f"Page {r['page_num']}: {clean_text}...")
                result = "\n".join(hints) if hints else "No results found for search."
                
            elif action == "GET_PAGE":
                try:
                    pnum = int(re.search(r'\d+', arg).group(0))
                    page = next((p for p in self.page_index.pages if p["page_num"] == pnum), None)
                    if page:
                        pages_seen.add(pnum)
                        tbl = page.get("tables", [])
                        tbl_str = "\n".join(tbl) if tbl else "None"
                        result = f"Page {pnum} Context:\n{page['text']}\nTables:\n{tbl_str}"
                    else:
                        result = f"Page {pnum} not found."
                except:
                    result = "Invalid page number. Provide an integer."
                    
            elif action == "GET_METADATA":
                import json
                sections = list(self.page_index.index.keys())
                result = f"Sections: {sections}\nTable of Contents Summary: \n{json.dumps(self.page_index.toc[:10], indent=2)}"
                
            else:
                result = f"Unknown action: {action}"
                
            memory.append({"step": step, "action": action, "arg": arg, "result": str(result)})
            
        return "Max agentic steps reached without confirmation."


def _fallback_generate_agent(llm, prompt: str) -> str:
    """Helper for agent generation calls."""
    if hasattr(llm, "generate"):
        return llm.generate(prompt)
    if hasattr(llm, "chat"):
        return llm.chat(prompt)
    if hasattr(llm, "ask"):
        return llm.ask(prompt)
    return ""

    def _verify_and_return(self, objective: str, final_answer: str) -> str:
        prompt = f"""Evaluate this final extracted answer against the objective.
Objective: {objective}
Extracted Data: {final_answer}

If the data answers the objective cleanly and reasonably, return exactly: OK
If it fundamentally misses the point or hallucinated heavily, return: REJECTED - <reason>
"""
        try:
            if hasattr(self.llm, "generate"):
                verdict = self.llm.generate(prompt)
            else:
                verdict = "OK"
        except:
            verdict = "OK"
            
        if verdict.strip().startswith("OK"):
            return final_answer
        return f"[Agentic Verification Warning: {verdict}]\nFallback Answer: {final_answer}"


def get_reasoning_retriever(file_path: str) -> ReasoningRetriever:
    return ReasoningRetriever(file_path)